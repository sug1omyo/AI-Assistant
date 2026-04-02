"""
OCR Integration for Chatbot
Provides text extraction from images and documents via Vision APIs.
Uses cloud Vision APIs (OpenAI, Gemini, Grok) - no local OCR service needed.
"""

import os
import base64
import requests
import logging
from pathlib import Path
from typing import Dict, Any, Tuple
import io

logger = logging.getLogger(__name__)


class OCRIntegration:
    """Handles OCR processing for uploaded files via Vision APIs."""

    def __init__(self):
        self.enabled = True

    @staticmethod
    def _detect_media_type(filename: str) -> str:
        ext = Path(filename).suffix.lower()
        return {
            '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
            '.gif': 'image/gif', '.webp': 'image/webp', '.bmp': 'image/bmp',
        }.get(ext, 'image/png')

    def extract_text_from_image(self, image_data: bytes, filename: str = "image.png") -> Dict[str, Any]:
        """Extract text from image using Vision APIs. Priority: OpenAI > Gemini > Grok."""
        result = {"success": False, "text": "", "confidence": 0, "language": "unknown", "method": "none"}

        base64_image = base64.b64encode(image_data).decode('utf-8')
        media_type = self._detect_media_type(filename)

        ocr_prompt = "Extract ALL text from this image. Return ONLY the extracted text, preserving layout where possible. If there's no text, return 'NO_TEXT_FOUND'."

        # Try OpenAI Vision
        openai_key = os.getenv("OPENAI_API_KEY")
        if openai_key:
            try:
                return self._ocr_openai(base64_image, media_type, filename, openai_key, ocr_prompt)
            except Exception as e:
                logger.warning(f"[OCR] OpenAI Vision failed: {e}")

        # Try Gemini Vision
        gemini_key = os.getenv("GEMINI_API_KEY_1") or os.getenv("GEMINI_API_KEY_2")
        if gemini_key:
            try:
                return self._ocr_gemini(base64_image, media_type, filename, gemini_key, ocr_prompt)
            except Exception as e:
                logger.warning(f"[OCR] Gemini Vision failed: {e}")

        # Try Grok Vision
        grok_key = os.getenv("GROK_API_KEY")
        if grok_key:
            try:
                return self._ocr_grok(base64_image, media_type, filename, grok_key, ocr_prompt)
            except Exception as e:
                logger.warning(f"[OCR] Grok Vision failed: {e}")

        result["error"] = "No Vision API key available for OCR"
        return result

    def _ocr_openai(self, base64_image: str, media_type: str, filename: str, api_key: str, prompt: str) -> Dict[str, Any]:
        """OCR via OpenAI Vision API (gpt-4o-mini)."""
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": "gpt-4o-mini",
                "messages": [{"role": "user", "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{base64_image}"}}
                ]}],
                "max_tokens": 4000
            },
            timeout=30
        )
        response.raise_for_status()
        text = response.json()["choices"][0]["message"]["content"].strip()
        if text and text != "NO_TEXT_FOUND":
            logger.info(f"[OCR] OpenAI extracted {len(text)} chars from {filename}")
            return {"success": True, "text": text, "confidence": 0.9, "language": "auto", "method": "openai_vision"}
        return {"success": False, "text": "", "confidence": 0, "language": "unknown", "method": "openai_vision"}

    def _ocr_gemini(self, base64_image: str, media_type: str, filename: str, api_key: str, prompt: str) -> Dict[str, Any]:
        """OCR via Gemini Vision API."""
        response = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}",
            headers={"Content-Type": "application/json"},
            json={"contents": [{"parts": [
                {"text": prompt},
                {"inline_data": {"mime_type": media_type, "data": base64_image}}
            ]}]},
            timeout=30
        )
        response.raise_for_status()
        data = response.json()
        text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
        if text and text != "NO_TEXT_FOUND":
            logger.info(f"[OCR] Gemini extracted {len(text)} chars from {filename}")
            return {"success": True, "text": text, "confidence": 0.88, "language": "auto", "method": "gemini_vision"}
        return {"success": False, "text": "", "confidence": 0, "language": "unknown", "method": "gemini_vision"}

    def _ocr_grok(self, base64_image: str, media_type: str, filename: str, api_key: str, prompt: str) -> Dict[str, Any]:
        """OCR via Grok/xAI Vision API (grok-4)."""
        response = requests.post(
            "https://api.x.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": "grok-4-1-fast-non-reasoning",
                "messages": [{"role": "user", "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{base64_image}"}}
                ]}],
                "max_tokens": 4000
            },
            timeout=30
        )
        response.raise_for_status()
        text = response.json()["choices"][0]["message"]["content"].strip()
        if text and text != "NO_TEXT_FOUND":
            logger.info(f"[OCR] Grok extracted {len(text)} chars from {filename}")
            return {"success": True, "text": text, "confidence": 0.85, "language": "auto", "method": "grok_vision"}
        return {"success": False, "text": "", "confidence": 0, "language": "unknown", "method": "grok_vision"}

    def extract_text_from_pdf(self, pdf_data: bytes, filename: str = "document.pdf") -> Dict[str, Any]:
        """Extract text from PDF with multi-library fallback + Vision OCR for scanned PDFs."""
        result = {"success": False, "text": "", "pages": 0, "method": "none"}

        # Try 1: pdfplumber (best for tables & complex layouts)
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(pdf_data)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    # Also try extracting tables
                    if not text.strip():
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                text += " | ".join(str(cell or "") for cell in row) + "\n"
                    text_parts.append(text)
                combined = "\n\n".join(text_parts).strip()
                if combined:
                    result["success"] = True
                    result["text"] = combined
                    result["pages"] = len(pdf.pages)
                    result["method"] = "pdfplumber"
                    logger.info(f"[OCR] pdfplumber extracted {len(combined)} chars from {filename}")
                    return result
                result["pages"] = len(pdf.pages)
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"[OCR] pdfplumber failed for {filename}: {e}")

        # Try 2: pypdf / PyPDF2
        for lib_name, import_path in [("pypdf", "pypdf"), ("PyPDF2", "PyPDF2")]:
            try:
                mod = __import__(import_path)
                reader = mod.PdfReader(io.BytesIO(pdf_data))
                text_parts = [page.extract_text() or "" for page in reader.pages]
                combined = "\n\n".join(text_parts).strip()
                if combined:
                    result["success"] = True
                    result["text"] = combined
                    result["pages"] = len(reader.pages)
                    result["method"] = lib_name
                    logger.info(f"[OCR] {lib_name} extracted {len(combined)} chars from {filename}")
                    return result
                result["pages"] = len(reader.pages)
                break  # Library worked but no text → likely scanned PDF
            except ImportError:
                continue
            except Exception as e:
                logger.warning(f"[OCR] {lib_name} failed for {filename}: {e}")

        # Try 3: Vision OCR fallback for scanned / image-based PDFs
        page_count = result.get("pages", 0) or 1
        logger.info(f"[OCR] Text extraction empty for {filename} ({page_count} pages), trying Vision OCR...")
        try:
            ocr_texts = self._ocr_pdf_pages(pdf_data, filename, max_pages=10)
            if ocr_texts:
                combined = "\n\n---\n\n".join(ocr_texts).strip()
                if combined:
                    result["success"] = True
                    result["text"] = combined
                    result["pages"] = len(ocr_texts)
                    result["method"] = "vision_ocr"
                    logger.info(f"[OCR] Vision OCR extracted {len(combined)} chars from {filename}")
                    return result
        except Exception as e:
            logger.warning(f"[OCR] Vision OCR fallback failed for {filename}: {e}")

        logger.error(f"[OCR] All PDF extraction methods failed for {filename}")
        return result

    def _ocr_pdf_pages(self, pdf_data: bytes, filename: str, max_pages: int = 10) -> list:
        """Convert PDF pages to images and OCR them via Vision API."""
        texts = []
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(pdf_data, first_page=1, last_page=max_pages, dpi=200)
        except ImportError:
            logger.warning("[OCR] pdf2image not installed, cannot OCR scanned PDF")
            return texts
        except Exception as e:
            logger.warning(f"[OCR] pdf2image conversion failed: {e}")
            return texts

        for i, img in enumerate(images):
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            img_bytes = buf.getvalue()
            result = self.extract_text_from_image(img_bytes, f"{filename}_page{i+1}.png")
            if result.get("success") and result.get("text", "").strip():
                texts.append(f"**Page {i+1}:**\n{result['text']}")
            else:
                texts.append(f"**Page {i+1}:** (no text extracted)")
        return texts

    def process_file(self, file_data: bytes, filename: str, content_type: str = None) -> Dict[str, Any]:
        """Process any file and extract text."""
        ext = Path(filename).suffix.lower()

        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.tiff']:
            return self.extract_text_from_image(file_data, filename)
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_data, filename)
        if ext in ['.txt', '.md', '.py', '.js', '.ts', '.html', '.css', '.json', '.xml',
                   '.csv', '.yaml', '.yml', '.toml', '.ini', '.cfg', '.log', '.sql', '.sh', '.bat']:
            try:
                text = file_data.decode('utf-8')
            except UnicodeDecodeError:
                text = file_data.decode('latin-1', errors='replace')
            return {"success": True, "text": text, "method": "direct_read", "file_type": ext[1:]}
        if ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_data, filename)
        if ext in ['.xlsx', '.xls']:
            return self._extract_from_excel(file_data, filename)
        return {"success": False, "text": "", "error": f"Unsupported file type: {ext}"}

    def _extract_from_docx(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from Word documents."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_data))
            text_parts = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    text_parts.append(" | ".join(cell.text for cell in row.cells))
            return {"success": True, "text": "\n".join(text_parts), "method": "python_docx", "file_type": "docx"}
        except Exception as e:
            logger.error(f"[OCR] DOCX extraction failed: {e}")
            return {"success": False, "text": "", "error": str(e)}

    def _extract_from_excel(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from Excel files."""
        try:
            import pandas as pd
            xlsx = pd.ExcelFile(io.BytesIO(file_data))
            text_parts = []
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                text_parts.append(df.to_string())
            return {"success": True, "text": "\n\n".join(text_parts), "method": "pandas",
                    "file_type": "xlsx", "sheets": xlsx.sheet_names}
        except Exception as e:
            logger.error(f"[OCR] Excel extraction failed: {e}")
            return {"success": False, "text": "", "error": str(e)}


# Global instance
ocr_client = OCRIntegration()


def extract_file_content(file_data: bytes, filename: str) -> Tuple[bool, str]:
    """Convenience function to extract content from file."""
    result = ocr_client.process_file(file_data, filename)
    return result.get("success", False), result.get("text", "")
